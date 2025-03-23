import re
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
from io import BytesIO

PUNCTUATION_KEYS = [
    "apostrophes", "colons", "commas", "curly_brackets", "double_inverted_commas",
    "ellipses", "em_dashes", "en_dashes", "exclamation_marks", "full_stops",
    "hyphens", "other_punctuation_marks", "question_marks", "round_brackets",
    "semicolons", "slashes", "square_brackets", "vertical_bars"
]

def extract_text(file):
    doc = Document(file)
    content = " ".join(p.text for p in doc.paragraphs)
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part:
                content += " ".join(p.text for p in part.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content += cell.text + " "
    content = re.sub(r'(\s?\.\s?){2,}', '...', content)
    return content

def count_punctuation(content):
    return {
        "apostrophes": len(re.findall(r"[\'’]", content)),
        "colons": len(re.findall(r":", content)),
        "commas": len(re.findall(r",", content)),
        "curly_brackets": len(re.findall(r"\{|\}", content)),
        "double_inverted_commas": len(re.findall(r"[“”\"]", content)),
        "ellipses": len(re.findall(r"…", content)) + content.count("..."),
        "em_dashes": len(re.findall(r"—", content)),
        "en_dashes": len(re.findall(r"–", content)),
        "exclamation_marks": len(re.findall(r"!", content)),
        "full_stops": len(re.findall(r"\.", content)) - content.count("..."),
        "hyphens": len(re.findall(r"-", content)),
        "other_punctuation_marks": len(re.findall(r"[*&%$@]", content)),
        "question_marks": len(re.findall(r"\?", content)),
        "round_brackets": len(re.findall(r"\(|\)", content)),
        "semicolons": len(re.findall(r";", content)),
        "slashes": len(re.findall(r"/", content)),
        "square_brackets": len(re.findall(r"\[|\]", content)),
        "vertical_bars": len(re.findall(r"\|", content)),
    }

def count_words(content):
    return len(re.findall(r"\b\w+\b", content))

def generate_graph(df, selected_keys):
    titles = df["filename"].str.replace(".docx", "", regex=False)
    fig, ax = plt.subplots(figsize=(12, 6))
    for key in selected_keys:
        ax.plot(titles, df[key], marker='o', label=key.replace('_', ' '))
    ax.set_title("Punctuation Frequency Across Documents")
    ax.set_xlabel("Document")
    ax.set_ylabel("Count")
    ax.legend()
    ax.grid(True)
    plt.xticks(rotation=45, ha='right')
    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf
